import os
import re
import xml.etree.ElementTree as ET
from docx import Document

# Define namespaces
namespaces = {
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'wp': 'http://schemas.openxmlformats.org/wordprocessingml/2006/wordprocessingDrawing'
}

months_map = {
    'january': '01', 'february': '02', 'march': '03', 'april': '04',
    'may': '05', 'june': '06', 'july': '07', 'august': '08',
    'september': '09', 'october': '10', 'november': '11', 'december': '12',
    'jan': '01', 'feb': '02', 'mar': '03', 'apr': '04', 'jun': '06', 'jul': '07', 'aug': '08', 'sep': '09', 'oct': '10', 'nov': '11', 'dec': '12'
}

def parse_date_from_doc(doc):
    for p in doc.paragraphs:
        p_text_lower = p.text.strip().lower()
        if p_text_lower.startswith('date:'):
            date_part = p_text_lower.replace('date:', '').strip()
            # Try pattern: Month Day(suffix), Year (e.g. "june 29th, 2025")
            m = re.search(r'([a-z]+)\s+(\d+)(?:st|nd|rd|th)?\s*,?\s*(\d{4})', date_part)
            if m:
                month_name, day, year = m.groups()
                month = months_map.get(month_name)
                if month:
                    return f"{year}-{month}-{int(day):02d}"
                    
            # Try pattern: Day(suffix) Month, Year (e.g. "29th june, 2025")
            m = re.search(r'(\d+)(?:st|nd|rd|th)?\s+([a-z]+)\s*,?\s*(\d{4})', date_part)
            if m:
                day, month_name, year = m.groups()
                month = months_map.get(month_name)
                if month:
                    return f"{year}-{month}-{int(day):02d}"
    return None

def clean_name(filename):
    # Remove extension
    name = os.path.splitext(filename)[0]
    
    # Strip user metadata prefixes
    name = re.sub(r'^(kibet brian|kibet_brian|190814|lab|–|\s|-)+', '', name, flags=re.IGNORECASE)
    
    # Create Title
    title = name.strip()
    # Normalize dashes and spaces
    title = re.sub(r'[\s_—–-]+', ' ', title)
    title = title.strip()
    
    # Capitalize appropriately (Title Case)
    words = title.split()
    capitalized_words = []
    for i, w in enumerate(words):
        if w.lower() in ['in', 'using', 'with', 'for', 'a', 'an', 'the', 'and', 'but', 'or', 'on', 'to', 'at', 'by', 'of'] and i > 0 and i < len(words) - 1:
            capitalized_words.append(w.lower())
        else:
            # Keep acronyms capitalized
            if w.isupper() and len(w) > 1:
                capitalized_words.append(w)
            elif w.lower() in ['w3af', 'ddos', 'hoic', 'nessus', 'openvas', 'zap', 'pmkid', 'bettercap', 'dvwa', 'sql', 'syn', 'nrpe', 'aws', 'is']:
                capitalized_words.append(w.upper())
                # Handle special capitalized acronyms
                if w.lower() == 'w3af': capitalized_words[-1] = 'w3af'
                if w.lower() == 'openvas': capitalized_words[-1] = 'OpenVAS'
                if w.lower() == 'bettercap': capitalized_words[-1] = 'Bettercap'
                if w.lower() == 'kali': capitalized_words[-1] = 'Kali'
                if w.lower() == 'nagios': capitalized_words[-1] = 'Nagios'
                if w.lower() == 'terraform': capitalized_words[-1] = 'Terraform'
            else:
                capitalized_words.append(w.capitalize())
                
    title = " ".join(capitalized_words)
    
    # Create Slug
    slug = title.lower()
    slug = re.sub(r'[^a-z0-9\s_-]', '', slug)
    slug = re.sub(r'[\s_]+', '-', slug)
    slug = re.sub(r'-+', '-', slug)
    slug = slug.strip('-')
    
    return title, slug

def extract_images_from_run(run, doc, slug, image_dir, image_count):
    images_in_run = []
    xml_str = run._r.xml
    try:
        root = ET.fromstring(xml_str)
    except Exception:
        return images_in_run, image_count

    blips = root.findall('.//a:blip', namespaces)
    for blip in blips:
        embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
        if embed_id and embed_id in doc.part.related_parts:
            image_part = doc.part.related_parts[embed_id]
            orig_ext = os.path.splitext(image_part.partname)[1]
            if not orig_ext:
                orig_ext = '.png'
            
            image_count += 1
            image_filename = f"image_{image_count}{orig_ext}"
            image_path = os.path.join(image_dir, image_filename)
            
            # Save the image binary data
            with open(image_path, 'wb') as f:
                f.write(image_part.blob)
                
            # Images are always mapped to projects directory since only projects contain images
            web_path = f"/images/projects/{slug}/{image_filename}"
            images_in_run.append(web_path)
            
    return images_in_run, image_count

def convert_docx_to_markdown(docx_path):
    filename = os.path.basename(docx_path)
    title, slug = clean_name(filename)
    
    # Load Document
    doc = Document(docx_path)
    
    # Extract date directly from the document content, with 2025-07-10 as the fallback published date
    extracted_date = parse_date_from_doc(doc)
    date_str = extracted_date if extracted_date else "2025-07-10"
    
    # Setup image directory (always projects since images = projects)
    image_dir = os.path.join("images", "projects", slug)
    os.makedirs(image_dir, exist_ok=True)
    
    # Define student/lecturer metadata prefixes to remove
    metadata_prefixes = [
        "student’s name:", "students name:", "student name:", "student's name:",
        "admission number:", "admission no:", "adm. no:", "adm no:",
        "lecturer’s name:", "lecturers name:", "lecturer name:", "lecturer’s name:", "lecturer's name:", "lecturer:",
        "course unit:", "course code:", "course:",
        "date:"
    ]
    
    # Find first regular text paragraph for excerpt
    first_text = ""
    for p in doc.paragraphs:
        p_text = p.text.strip()
        if not p_text:
            continue
        if p.style.name.startswith('Heading'):
            continue
        p_text_lower = p_text.lower()
        if any(p_text_lower.startswith(prefix) for prefix in metadata_prefixes):
            continue
        first_text = p_text
        break
        
    excerpt = re.sub(r'[\*\`_\n"]', ' ', first_text)
    excerpt = re.sub(r'\s+', ' ', excerpt).strip()
    if len(excerpt) > 180:
        excerpt = excerpt[:177] + "..."
    if not excerpt:
        excerpt = f"Cybersecurity post for {title}."
        
    body_lines = []
    image_count = 0
    
    # Parse elements
    for p in doc.paragraphs:
        p_text_stripped = p.text.strip()
        p_text_lower = p_text_stripped.lower()
        
        # Skip metadata lines
        if any(p_text_lower.startswith(prefix) for prefix in metadata_prefixes):
            continue
            
        p_text_parts = []
        
        # Determine heading level
        heading_level = 0
        if p.style.name.startswith('Heading'):
            try:
                heading_level = int(p.style.name.replace('Heading', '').strip())
            except ValueError:
                pass
                
        for run in p.runs:
            run_text = run.text
            
            # Process formatting
            if run.bold and run_text.strip():
                lead_space = run_text[:len(run_text)-len(run_text.lstrip())]
                trail_space = run_text[len(run_text.rstrip()):]
                core_text = run_text.strip()
                run_text = f"{lead_space}**{core_text}**{trail_space}"
            if run.italic and run_text.strip():
                lead_space = run_text[:len(run_text)-len(run_text.lstrip())]
                trail_space = run_text[len(run_text.rstrip()):]
                core_text = run_text.strip()
                run_text = f"{lead_space}*{core_text}*{trail_space}"
                
            # Extract images in run
            run_images, image_count = extract_images_from_run(run, doc, slug, image_dir, image_count)
            for img_url in run_images:
                run_text += f"\n\n![Screenshot]({img_url})\n\n"
                
            p_text_parts.append(run_text)
            
        paragraph_text = "".join(p_text_parts)
        paragraph_text = re.sub(r'\n{3,}', '\n\n', paragraph_text)
        
        # Format the line
        formatted_line = ""
        if heading_level > 0:
            prefix = "#" * heading_level + " "
            formatted_line = prefix + paragraph_text.strip()
        elif p.style.name == 'List Bullet':
            formatted_line = f"* {paragraph_text.strip()}"
        elif p.style.name == 'List Number':
            formatted_line = f"1. {paragraph_text.strip()}"
        else:
            if paragraph_text.strip():
                formatted_line = paragraph_text.strip()
                
        if formatted_line:
            body_lines.append(formatted_line)
            
    # Deduplicate consecutive lines (e.g. duplicated headings)
    def strip_markdown(text):
        t = re.sub(r'[\*\`_\n"]', '', text)
        t = re.sub(r'^#+\s+', '', t)
        return t.strip().lower()

    filtered_body_lines = []
    for line in body_lines:
        if not line.strip():
            filtered_body_lines.append(line)
            continue
        if not filtered_body_lines:
            filtered_body_lines.append(line)
            continue
            
        # Find last non-empty line
        last_idx = -1
        for idx in range(len(filtered_body_lines) - 1, -1, -1):
            if filtered_body_lines[idx].strip():
                last_idx = idx
                break
                
        if last_idx != -1:
            last_line = filtered_body_lines[last_idx]
            clean_last = strip_markdown(last_line)
            clean_curr = strip_markdown(line)
            
            if clean_last == clean_curr:
                last_is_formatted = last_line.startswith('#') or ('**' in last_line) or ('*' in last_line)
                curr_is_formatted = line.startswith('#') or ('**' in line) or ('*' in line)
                
                if curr_is_formatted and not last_is_formatted:
                    filtered_body_lines[last_idx] = line
                    continue
                elif last_is_formatted and not curr_is_formatted:
                    continue
                    
        filtered_body_lines.append(line)

    # Also parse tables if any
    table_lines = []
    for table in doc.tables:
        table_lines.append("")
        for i, row in enumerate(table.rows):
            row_cells = [cell.text.strip().replace('\n', ' <br> ') for cell in row.cells]
            table_lines.append("| " + " | ".join(row_cells) + " |")
            if i == 0:
                table_lines.append("| " + " | ".join(["---"] * len(row_cells)) + " |")
        table_lines.append("")

    # Determine collection based on final image count (screenshots/diagrams go to projects; text goes to writeups)
    if image_count > 0:
        collection = 'portfolio'
        output_dir = '_portfolio'
        output_filename = f"{slug}.md"
    else:
        collection = 'posts'
        output_dir = '_posts'
        output_filename = f"{date_str}-{slug}.md"

    # Assemble Markdown with correct front-matter headers
    markdown_lines = []
    markdown_lines.append("---")
    markdown_lines.append(f'title: "{title}"')
    if collection == 'portfolio':
        markdown_lines.append("collection: portfolio")
        markdown_lines.append(f"permalink: /portfolio/{slug}/")
    else:
        markdown_lines.append(f"date: {date_str}")
        markdown_lines.append(f"permalink: /posts/{date_str.replace('-', '/')}/{slug}/")
    markdown_lines.append(f'excerpt: "{excerpt}"')
    markdown_lines.append("tags:")
    if collection == 'portfolio':
        markdown_lines.append("  - project")
        markdown_lines.append("  - cybersecurity")
    else:
        markdown_lines.append("  - writeup")
        markdown_lines.append("  - reflection")
        markdown_lines.append("  - cybersecurity")
    markdown_lines.append("---")
    markdown_lines.append("")
    
    # Write filtered lines to markdown
    for line in filtered_body_lines:
        markdown_lines.append(line)
        markdown_lines.append("")
        
    markdown_lines.extend(table_lines)

    markdown_content = "\n".join(markdown_lines)
    markdown_content = re.sub(r'\n{3,}', '\n\n', markdown_content)
    
    # Save markdown file
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, output_filename)
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(markdown_content)
        
    # Clean up empty image directory if no images were extracted
    if image_count == 0 and os.path.exists(image_dir):
        try:
            os.rmdir(image_dir)
        except OSError:
            pass
            
    print(f"Converted: '{filename}' -> '{output_dir}/{output_filename}' (Extracted {image_count} images)")

def main():
    # Scan both folders to find all source docx files
    folders = ["pending_writeups", "pending_projects"]
    processed_files = set()
    
    for folder in folders:
        if os.path.exists(folder):
            for filename in os.listdir(folder):
                if filename.endswith(".docx") and not filename.startswith("~$"):
                    if filename in processed_files:
                        continue
                    processed_files.add(filename)
                    
                    docx_path = os.path.join(folder, filename)
                    try:
                        convert_docx_to_markdown(docx_path)
                    except Exception as e:
                        print(f"Error converting '{filename}': {e}")

if __name__ == "__main__":
    main()
